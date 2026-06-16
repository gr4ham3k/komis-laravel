# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'

# Title
title = doc.add_heading('MotoKomis - Powtorzenie kodu', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Full-stack Laravel + Blade aplikacja do ogloszen motoryzacyjnych i uslug')

doc.add_page_break()

# ============================================================
# 1. STRUKTURA PROJEKTU
# ============================================================
doc.add_heading('1. Struktura projektu', level=1)
doc.add_paragraph(
    'Aplikacja MotoKomis zbudowana jest w architekturze MVC (Model-View-Controller) '
    'na frameworku Laravel 13.15.0 z PHP 8.4 i PostgreSQL. '
    'Widoki renderowane sa po stronie serwera za pomoca silnika Blade.'
)

doc.add_heading('Glowne katalogi:', level=2)
items = [
    'app/Http/Controllers/ - kontrolery logiki biznesowej',
    'app/Models/ - modele Eloquent (Listing, Brand, CarModel, itd.)',
    'resources/views/ - widoki Blade',
    'routes/web.php - definicje routeow',
    'database/migrations/ - migracje bazy danych',
    'database/seeders/ - seedery (dane przykladowe)',
]
for i in items:
    doc.add_paragraph(i, style='List Bullet')

doc.add_page_break()

# ============================================================
# 2. STRONA GLOWNA (home)
# ============================================================
doc.add_heading('2. Strona glowna (Home)', level=1)

doc.add_heading('2.1. Route', level=2)
doc.add_paragraph('routes/web.php:1')
p = doc.add_paragraph()
p.style = doc.styles['Intense Quote']
run = p.add_run("Route::get('/', [ListingController::class, 'home'])->name('home');")

doc.add_heading('2.2. Kontroler - ListingController::home()', level=2)
doc.add_paragraph('app/Http/Controllers/ListingController.php:9-32')
doc.add_paragraph(
    'Metoda home() wykonuje trzy zapytania do bazy:\n'
    '  - Pobiera 6 ostatnich aktywnych ogloszen (z relacjami: brand, carModel, fuel, transmission, bodyType, tags, images)\n'
    '  - Pobiera 8 najpopularniejszych tagow (liczone po ilosci aktywnych ogloszen)\n'
    '  - Zbiera statystyki: liczba aktywnych ogloszen, marek i uslug\n'
    'Nastepnie zwraca widok home.blade.php z tymi danymi.'
)

doc.add_heading('2.3. Widok - home.blade.php', level=2)
doc.add_paragraph('resources/views/home.blade.php')
doc.add_paragraph(
    'Strona glowna sklada sie z:\n'
    '  - Hero section - tlo z gradientem i zdjeciem BMW, tekst powitalny, dwa przyciski (ogloszenia, uslugi)\n'
    '  - Panel szybkiego wyszukiwania - formularz z polem "q" przekierowuje do /listings?q=...\n'
    '  - Statystyki - 3 kafelki: aktywne ogloszenia, marki w bazie, aktywne uslugi\n'
    '  - Popularne tagi - lista szybkich tagow (#Sedan, #Klimatyzacja itp.) z licznikami\n'
    '  - Polecane ogloszenia - grid 6 najnowszych ogloszen (renderowane przez _card.blade.php)\n\n'
    'Layout: rozszerza layouts.app - naglowek z navbarem i stopke.'
)

doc.add_page_break()

# ============================================================
# 3. STRONA OGLOSZEN (listings/index)
# ============================================================
doc.add_heading('3. Lista ogloszen (Listings index)', level=1)

doc.add_heading('3.1. Route', level=2)
p = doc.add_paragraph()
p.style = doc.styles['Intense Quote']
run = p.add_run("Route::get('/listings', [ListingController::class, 'index'])->name('listings.index');")

doc.add_heading('3.2. Kontroler - ListingController::index()', level=2)
doc.add_paragraph('app/Http/Controllers/ListingController.php:34-84')
doc.add_paragraph(
    'Metoda index() buduje zapytanie z wieloma filtrami:\n\n'
    '  - Wyszukiwanie tekstowe (q) - uzywa pg_trgm similarity do wyszukiwania po tytule, opisie, miescie, '
    'kolorze, marce, modelu\n'
    '  - Filtry dokladne: brand_id, model_id, fuel_id, transmission_id, body_type_id\n'
    '  - Lokalizacja: filtrowanie po wspolrzednych (lat/lng) z promieniem (km). '
    'Jesli nie ma wspolrzednych, filtruje po nazwie miasta (city)\n'
    '  - Zakresy: price_min/max, year_min/max, mileage_max, power_min, engine_min\n'
    '  - Tagi: AND - ogloszenie musi miec WSZYSTKIE zaznaczone tagi\n'
    '  - Sortowanie: price_asc, price_desc, year_desc, mileage_asc, popular (views_count), domyslnie latest\n'
    '  - Paginacja: per_page = 10, 20, 50 (domyslnie 12)\n\n'
    'Wynik: $listings->paginate($perPage)->withQueryString() - przekazuje do widoku listings.index'
)

doc.add_heading('3.3. Widok - index.blade.php', level=2)
doc.add_paragraph('resources/views/listings/index.blade.php')
doc.add_paragraph(
    'Uklad dwukolumnowy:\n'
    '  - Lewa kolumna (col-lg-3): filtr boczny (_filters.blade.php)\n'
    '  - Prawa kolumna (col-lg-9): lista ogloszen\n\n'
    'Gorny pasek: liczba znalezionych ogloszen + sortowanie (select) + per_page (select)\n'
    'Grid ogloszen: 3 kolumny na desktopie, kazda karta to _card.blade.php\n'
    'Paginacja: wysrodkowana na dole ($listings->links())\n\n'
    'Porownanie: przycisk u gory, licznik w badge. Kazda karta ma przycisk "Dodaj do porownania" / "Usun z porownania".'
)

doc.add_heading('3.4. Filtr boczny - _filters.blade.php', level=2)
doc.add_paragraph('resources/views/listings/_filters.blade.php')
doc.add_paragraph(
    'Formularz GET z nastepujacymi polami:\n'
    '  - q - wyszukiwarka tekstowa\n'
    '  - Marka (brand_id) - autocomplete AJAX przez /brands/search z 200ms debounce\n'
    '  - Model (model_id) - autocomplete AJAX przez /models/search, wymaga wybranej marki\n'
    '  - Cena od/do (price_min, price_max)\n'
    '  - Rok produkcji od/do (year_min, year_max)\n'
    '  - Przebieg do (mileage_max)\n'
    '  - Lokalizacja - mapa Leaflet z mozliwoscia klikniecia (lat/lng) + promien (radius)\n'
    '  - Szczegoly techniczne (accordion): paliwo, skrzynia biegow, nadwozie, moc min, pojemnosc min\n'
    '  - Wyposazenie (accordion): checkboxy tagow\n\n'
    'JavaScript: autocomplete fetch z onmousedown (zamiast onclick, aby nie konfliktowac z blurem). '
    'Mapa: Leaflet z OpenStreetMap, mozliwosc klikniecia, ustawienia markera i kola promienia.'
)

doc.add_heading('3.5. Karta ogloszenia - _card.blade.php', level=2)
doc.add_paragraph('resources/views/listings/_card.blade.php')
doc.add_paragraph(
    'Pojedyncza karta ogloszenia zawiera:\n'
    '  - Zdjecie (pierwsze z relacji images) lub ikone aparatu\n'
    '  - Tytul (link do show)\n'
    '  - Status (Aktywne/Nieaktywne) - badge\n'
    '  - Marka + Model + Miasto\n'
    '  - Cena (formatowana ze spacjami)\n'
    '  - Przycisk "Dodaj do porownania" / "Usun z porownania" (zaleznie od sesji)'
)

doc.add_page_break()

# ============================================================
# 4. PANEL ADMINA
# ============================================================
doc.add_heading('4. Panel admina', level=1)
doc.add_paragraph(
    'Wszystkie route admina sa w grupie middleware [auth, admin]. '
    'Middleware "admin" sprawdza, czy zalogowany uzytkownik ma flage is_admin = true.'
)

doc.add_heading('4.1. Slowniki (admin/dictionaries)', level=2)
doc.add_heading('Route:', level=3)
p = doc.add_paragraph()
p.style = doc.styles['Intense Quote']
run = p.add_run("Route::prefix('admin/dictionaries') -> DictionaryController@index")
doc.add_paragraph(
    'CRUD dla slownikow: brandow, modeli, paliw, skrzyn biegow, nadwozi, tagow. '
    'Kazdy ma osobny kontroler w Admin/Dictionary/ '
    '(BrandController, ModelController, FuelController, TransmissionController, BodyTypeController, TagController) '
    'z metodami store(), update(), destroy().'
)

doc.add_heading('Kontroler - DictionaryController@index()', level=3)
doc.add_paragraph('app/Http/Controllers/Admin/DictionaryController.php')
doc.add_paragraph(
    'Pobiera wszystkie rekordy z tabel: brands, car_models, fuels, transmissions, body_types, tags '
    'i przekazuje do widoku admin.dictionaries.'
)

doc.add_heading('Widok - dictionaries.blade.php', level=3)
doc.add_paragraph('resources/views/admin/dictionaries.blade.php')
doc.add_paragraph(
    'Zakladki (Bootstrap tabs) dla kazdego slownika. W kazdej zakladce:\n'
    '  - Formularz dodawania (input + przycisk "Dodaj")\n'
    '  - Lista istniejacych rekordow z przyciskami Edytuj/Usun\n'
    '  - Edycja inline z zapisem przez AJAX/PATCH\n\n'
    'Na gorze pasek nawigacyjny z przyciskami: Uslugi, Uzytkownicy.'
)

doc.add_heading('4.2. Uslugi (admin/services)', level=2)
doc.add_heading('Kontroler - ServiceAdminController', level=3)
doc.add_paragraph('app/Http/Controllers/Admin/ServiceAdminController.php')
doc.add_paragraph(
    'Pelny CRUD dla uslug motoryzacyjnych:\n'
    '  - index() - lista wszystkich uslug\n'
    '  - store() - dodawanie\n'
    '  - update() - edycja\n'
    '  - destroy() - usuwanie'
)

doc.add_heading('Widok - services.blade.php', level=3)
doc.add_paragraph('resources/views/admin/services.blade.php')
doc.add_paragraph(
    'Tabela z uslugami (nazwa, opis, status, kategoria, uzytkownik). '
    'Formularz dodawania/edycji w modalu. Przyciski: Uzytkownicy, Slowniki.'
)

doc.add_heading('4.3. Uzytkownicy (admin/users)', level=2)
doc.add_heading('Kontroler - UserAdminController', level=3)
doc.add_paragraph('app/Http/Controllers/Admin/UserAdminController.php')
doc.add_paragraph(
    'Zarzadzanie uzytkownikami:\n'
    '  - index() - lista uzytkownikow z paginacja\n'
    '  - store() - tworzenie nowego uzytkownika\n'
    '  - update() - edycja (nazwa, email, admin)\n'
    '  - toggleBan() - banowanie/odbanowanie (ustawia banned_at)\n'
    '  - destroy() - usuwanie konta'
)

doc.add_heading('Widok - users.blade.php', level=3)
doc.add_paragraph('resources/views/admin/users.blade.php')
doc.add_paragraph(
    'Tabela z uzytkownikami: ID, nazwa, email, rola (admin/tak/nie), status (ban/tak/nie), '
    'liczba ogloszen, akcje (edytuj, ban/unban, usun). Modal dodawania/edycji.'
)

doc.add_page_break()

# ============================================================
# 5. PANEL UZYTKOWNIKA
# ============================================================
doc.add_heading('5. Panel uzytkownika', level=1)
doc.add_paragraph(
    'Route w grupie middleware auth. Panel dostepny po zalogowaniu.'
)

doc.add_heading('5.1. Dashboard (user/panel)', level=2)
doc.add_heading('Kontroler - UserPanelController@dashboard()', level=3)
doc.add_paragraph('app/Http/Controllers/UserPanelController.php:12-31')
doc.add_paragraph(
    'Zbiera statystyki zalogowanego uzytkownika:\n'
    '  - Liczba ogloszen (wszystkie + aktywne)\n'
    '  - Liczba uslug\n'
    '  - Liczba konwersacji (jako user2)\n'
    '  - Laczna liczba wyswietlen ogloszen\n'
    '  - 5 ostatnich ogloszen uzytkownika\n\n'
    'Zwraca widok user.dashboard.'
)

doc.add_heading('Widok - dashboard.blade.php', level=3)
doc.add_paragraph('resources/views/user/dashboard.blade.php')
doc.add_paragraph(
    'Naglowek powitalny z nazwa uzytkownika. Statystyki w kafelkach. '
    'Lista 5 ostatnich ogloszen z linkami do edycji. '
    'Boczna nawigacja: Moje ogloszenia, Moje uslugi, Konwersacje, Moj profil.'
)

doc.add_heading('5.2. Moje ogloszenia (user/listings)', level=2)
doc.add_heading('Kontroler - UserPanelController@listings()', level=3)
doc.add_paragraph('app/Http/Controllers/UserPanelController.php:33-41')
doc.add_paragraph(
    'Pobiera ogloszenia zalogowanego uzytkownika, paginacja 10 na strone.'
)

doc.add_heading('5.3. Edycja ogloszen (user/listings/{id}/edit)', level=2)
doc.add_heading('Kontroler - ListingController@edit() i @update()', level=3)
doc.add_paragraph(
    'edit() - sprawdza ownership (user_id == Auth::id()), 403 jesli nie. '
    'Laduje wszystkie slowniki i zwraca widok listings.edit.\n'
    'update() - taka sama walidacja jak store(), aktualizuje ogloszenie i synchonizuje tagi.'
)

doc.add_heading('5.4. Moj profil (user/profile)', level=2)
doc.add_heading('Kontroler - ProfileController', level=3)
doc.add_paragraph('app/Http/Controllers/ProfileController.php')
doc.add_paragraph(
    '  - edit() - zwraca widok user.profile\n'
    '  - update() - zmiana name/email (walidacja unique email z pominieciem biezacego uzytkownika)\n'
    '  - password() - zmiana hasla z walidacja current_password'
)

doc.add_heading('Widok - profile.blade.php', level=3)
doc.add_paragraph('resources/views/user/profile.blade.php')
doc.add_paragraph(
    'Dwie kolumny obok siebie:\n'
    '  - Lewa: formularz edycji danych (name, email)\n'
    '  - Prawa: formularz zmiany hasla (current_password, new_password, confirm)'
)

doc.add_page_break()

# ============================================================
# 6. DODATKOWE FUNKCJONALNOSCI
# ============================================================
doc.add_heading('6. Dodatkowe funkcjonalnosci', level=1)

doc.add_heading('6.1. Porownywanie ogloszen', level=2)
doc.add_paragraph(
    'Route: /compare/* (ListingCompareController)\n'
    'Dziala na sesji (compare_listings - tablica ID). '
    'Mozna dodawac, usuwac, czyscic. Widok compare/index.blade.php '
    'wyswietla tabele porownawcza z parametrami aut.'
)

doc.add_heading('6.2. Konwersacje/Czat', level=2)
doc.add_paragraph(
    'Route: /conversations/* (ConversationController) i /messages (MessageController)\n'
    'System wiadomosci miedzy uzytkownikami. Rozpoczecie czatu z /chat/start/{listingId} '
    'tworzy lub otwiera istniejaca konwersacje. Wiadomosci ladowane przez AJAX.'
)

doc.add_heading('6.3. Uslugi', level=2)
doc.add_paragraph(
    'Route: /services/* (ServiceController)\n'
    'Uzytkownicy moga dodawac uslugi motoryzacyjne (np. wulkanizacja, detailing). '
    'Przegladanie, dodawanie opinii (rating + komentarz).'
)

doc.add_heading('6.4. Autoryzacja', level=2)
doc.add_paragraph(
    'Login/rejestracja przez AuthController (routes/web.php:24-29). '
    'Logowanie przez POST /logout. Panel admina sprawdza middleware admin.'
)

doc.add_heading('6.5. API', level=2)
doc.add_paragraph(
    'routes/api.php - apiResource dla ListingApiController:\n'
    '  - GET /api/listings - lista aktywnych ogloszen (JSON)\n'
    '  - GET /api/listings/{id} - szczegoly ogloszenia (JSON)\n'
    '  - POST /api/listings - dodawanie (wymaga autoryzacji)\n'
    '  - PUT /api/listings/{id} - edycja\n'
    '  - DELETE /api/listings/{id} - usuwanie\n\n'
    'Uwaga: API nie ma jeszcze token-based auth (Sanctum).'
)

doc.add_page_break()

# ============================================================
# 7. WAZNE MECHANIZMY
# ============================================================
doc.add_heading('7. Wazne mechanizmy w kodzie', level=1)

doc.add_heading('7.1. Autocomplete AJAX dla marek/modeli', level=2)
doc.add_paragraph(
    '- endpoint /brands/search?q=... uzywa pg_trgm similarity (gdy query >= 3 znaki) lub LIKE (gdy < 3)\n'
    '- endpoint /models/search?brand_id=...&q=... - analogicznie, ale filtrowane po brand_id\n'
    '- Frontend: 200ms debounce, onmousedown zamiast onclick, blur z 300ms opoznieniem\n'
    '- Wybor zapisuje ID w hidden input, czysci drugie pole\n'
    '- Wyniki wyswietlane w absolute divie (list-group)'
)

doc.add_heading('7.2. Wyszukiwanie po lokalizacji', level=2)
doc.add_paragraph(
    '- Mapa Leaflet z OpenStreetMap\n'
    '- Klikniecie na mape ustawia marker + zapisuje lat/lng w hidden inputach\n'
    '- Wybor promienia (5-150 km lub cala Polska)\n'
    '- Obliczanie odleglosci w SQL: wzor Haversine (6371 km * acos(...))\n'
    '- Geocoding: wpisanie miasta w pole i klikniecie "Szukaj" ustawia marker'
)

doc.add_heading('7.3. pg_trgm (Trigram) - wyszukiwanie podobienstwa', level=2)
doc.add_paragraph(
    '- Baza PostgreSQL z rozszerzeniem pg_trgm\n'
    '- similarity() - zwraca podobienstwo 0-1\n'
    '- Operator % - "podobny do"\n'
    '- Uzywane w: wyszukiwarce tekstowej (q), autocomplete marek/modeli'
)

doc.add_heading('7.4. Policy-Based Authorization (Admin)', level=2)
doc.add_paragraph(
    '- is_admin - kolumna boolean w tabeli users\n'
    '- Admin middleware: sprawdza Auth::check() && Auth::user()->is_admin\n'
    '- Ownership: kazda metoda edycji/usuwania sprawdza user_id == Auth::id() lub 403'
)

doc.add_page_break()

# ============================================================
# 8. SCHEMAT ROUTINGU
# ============================================================
doc.add_heading('8. Schemat routingu (web.php)', level=1)

items = [
    ('GET /', 'home', 'Strona glowna'),
    ('GET /listings', 'listings.index', 'Lista ogloszen'),
    ('GET /listings/{listing}', 'listings.show', 'Szczegoly ogloszenia'),
    ('GET /listings/create', 'listings.create', 'Dodaj ogloszenie'),
    ('POST /listings/create', 'listings.store', 'Zapisz ogloszenie'),
    ('GET /user/panel', 'user.panel', 'Panel uzytkownika'),
    ('GET /user/listings', 'my.listings', 'Moje ogloszenia'),
    ('GET /user/profile', 'user.profile', 'Moj profil'),
    ('PUT /user/profile', 'user.profile.update', 'Aktualizacja profilu'),
    ('PUT /user/profile/password', 'user.profile.password', 'Zmiana hasla'),
    ('GET /admin/dictionaries', 'admin.dictionaries.index', 'Slowniki (admin)'),
    ('GET /admin/services', 'admin.services.index', 'Uslugi (admin)'),
    ('GET /admin/users', 'admin.users.index', 'Uzytkownicy (admin)'),
    ('GET /conversations', 'conversations.index', 'Konwersacje'),
    ('GET /services', 'services.index', 'Uslugi'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Metoda + URL'
hdr[1].text = 'Nazwa route'
hdr[2].text = 'Opis'

for method, name, desc in items:
    row = table.add_row().cells
    row[0].text = method
    row[1].text = name
    row[2].text = desc

doc.add_page_break()

# ============================================================
# 9. MODEL BAZY DANYCH (Listing)
# ============================================================
doc.add_heading('9. Model Listing - kluczowe pola', level=1)
doc.add_paragraph(
    'Tabela listings jest glowna tabela aplikacji. Kluczowe kolumny:\n\n'
    '  - title, description - tresc ogloszenia\n'
    '  - price (decimal) - cena w PLN\n'
    '  - city, latitude, longitude - lokalizacja\n'
    '  - year (integer) - rok produkcji\n'
    '  - brand_id, model_id - FK do slownikow\n'
    '  - fuel_id, transmission_id, body_type_id - FK do slownikow\n'
    '  - color, mileage, engine_capacity, power_hp - parametry techniczne\n'
    '  - user_id - FK do users (wlasciciel)\n'
    '  - status - active / inactive\n'
    '  - views_count - licznik wyswietlen\n'
    '  - Relacje: images (hasMany), tags (belongsToMany), brand, carModel, fuel, transmission, bodyType'
)

doc.add_paragraph('')
doc.add_paragraph('-- Koniec dokumentu --')

doc.save('MotoKomis_powtorzenie_kodu.docx')
print("DONE: MotoKomis_powtorzenie_kodu.docx created")
